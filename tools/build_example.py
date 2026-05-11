"""
Build example/bundle assets and example/report.zip (~10 pages in Word when generated).

  python tools/build_example.py

Requires: pandas, pillow (install: pip install -e ".[dev]")
"""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
BUNDLE = ROOT / "example" / "bundle"
ZIP_PATH = ROOT / "example" / "report.zip"


def _font(size: int = 22) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for name in ("arial.ttf", "DejaVuSans.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _long_text(title: str, *, segments: int = 6, repeats: int = 9) -> str:
    base = (
        "This fixed narrative ships inside the ZIP for pagination. It states controls, "
        "baselines, and trace references. Cross-reference other parts of the report using "
        "labels such as Table 1 or Figure 1 where relevant. "
    )
    chunks: list[str] = []
    for i in range(segments):
        chunks.append(f"{title} — segment {i + 1}. " + " ".join(base for _ in range(repeats)))
    return "\n\n".join(chunks)


def _write_csv(path: Path, stem: str, rows: int = 26) -> None:
    rows_data = []
    for i in range(rows):
        rows_data.append(
            {
                "id": i + 1,
                "workstream": f"WS-{(i % 5) + 1}",
                "metric": round(10.0 + (i % 7) * 1.1, 2),
                "status": ["open", "closed", "watch"][i % 3],
                "owner": f"O{(i % 4) + 1}",
                "note": f"{stem} row {i + 1}",
            }
        )
    pd.DataFrame(rows_data).to_csv(path, index=False)


def _fig(path: Path, title: str, rgb: tuple[int, int, int]) -> None:
    im = Image.new("RGB", (960, 560), (248, 250, 252))
    dr = ImageDraw.Draw(im)
    dr.rectangle((24, 24, 936, 536), outline=rgb, width=5)
    f = _font(24)
    dr.text((40, 40), title, fill=(17, 24, 39), font=f)
    dr.text((40, 280), "Example figure for docgen sample report.", fill=(55, 65, 81), font=_font(16))
    im.save(path)


def main() -> None:
    content = BUNDLE / "content"
    tables = BUNDLE / "tables"
    images = BUNDLE / "images"
    for d in (content, tables, images):
        d.mkdir(parents=True, exist_ok=True)

    # Word template with header/footer (used by docgen when ZIP has exactly one .docx)
    tmpl = BUNDLE / "template.docx"
    tdoc = Document()
    sec = tdoc.sections[0]
    sec.header.paragraphs[0].text = "Example header — docgen template"
    sec.footer.paragraphs[0].text = "Example footer — page { PAGE }"
    tdoc.save(str(tmpl))

    (content / "intro.txt").write_text(_long_text("Introduction"), encoding="utf-8")
    (content / "scope.txt").write_text(_long_text("Scope narrative"), encoding="utf-8")
    (content / "standards.txt").write_text(_long_text("Standards and controls"), encoding="utf-8")
    (content / "methodology.txt").write_text(_long_text("Methodology"), encoding="utf-8")
    (content / "discussion.txt").write_text(_long_text("Discussion"), encoding="utf-8")

    _write_csv(tables / "scope_items.csv", "scope", 26)
    _write_csv(tables / "stakeholders.csv", "stakeholder", 26)
    _write_csv(tables / "milestones.csv", "milestone", 26)
    _write_csv(tables / "program_inventory.csv", "program", 26)
    _write_csv(tables / "quality_checks.csv", "quality", 26)

    _fig(images / "fig_context.png", "System context (example)", (37, 99, 235))
    _fig(images / "fig_controls.png", "Control flow (example)", (180, 83, 9))
    _fig(images / "fig_metrics.png", "Metrics overview (example)", (22, 163, 74))
    _fig(images / "fig_summary.png", "Summary dashboard (example)", (124, 58, 237))

    payload = {
        "intro_llm": {
            "program": "Quarterly evidence pack",
            "audience": "Program steering group",
            "figure": "fig_context.png",
            "table": "scope_items.csv",
        },
        "objectives_llm": {
            "primary": "Complete evidence trail for gate review",
            "metric": "95% artifacts linked to requirements",
        },
        "scope_llm": {
            "in_scope": "Reporting, analytics extracts, governance slides",
            "out_of_scope": "Vendor contract negotiation",
        },
        "methodology_llm": {
            "approach": "Structured data pulls plus manual spot checks",
            "validation": "10% independent QA sample",
        },
        "results_llm": {
            "headline": "KPIs green; two actions tracked to closure",
            "figure": "fig_metrics.png",
            "table": "quality_checks.csv",
        },
        "discussion_llm": {
            "theme": "Delivery stable; monitor dependency WS-2",
        },
        "limitations_llm": {
            "data": "Extracts lag source systems by up to two business days",
            "coverage": "Self-reported status for some workstreams",
        },
        "conclusion_llm": {
            "verdict": "Ready for sign-off pending final attachment upload",
            "figure": "fig_summary.png",
        },
    }
    (BUNDLE / "bundle.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")

    ZIP_PATH.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in BUNDLE.rglob("*"):
            if p.is_file():
                zf.write(p, arcname=p.relative_to(BUNDLE).as_posix())

    print(f"Wrote {BUNDLE} and {ZIP_PATH}")


if __name__ == "__main__":
    main()
