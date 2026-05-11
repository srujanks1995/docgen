"""Bookmarks, internal hyperlinks, and page-reference fields for python-docx."""

from __future__ import annotations

from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


class BookmarkCounter:
    def __init__(self) -> None:
        self._n = 0

    def next_id(self) -> int:
        self._n += 1
        return self._n


def add_bookmark(paragraph: Paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    paragraph._element.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._element.append(end)


def add_hyperlink_to_bookmark(paragraph: Paragraph, text: str, bookmark_name: str, _doc_part) -> None:
    """Append an internal hyperlink (w:anchor) to the paragraph."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_pageref_field(paragraph: Paragraph, bookmark_name: str) -> None:
    """Append a PAGE field for the page where ``bookmark_name`` is located (updates in Word)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def append_toc_entry_with_pageref(
    paragraph: Paragraph,
    doc_part,
    link_display: str,
    bookmark_name: str,
    *,
    tab_stop_inches: float = 6.2,
) -> None:
    """Hyperlinked entry, dot leader, and right-aligned page number (PAGEREF)."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(tab_stop_inches),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS,
    )
    add_hyperlink_to_bookmark(paragraph, link_display, bookmark_name, doc_part)
    paragraph.add_run("\t")
    add_pageref_field(paragraph, bookmark_name)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph as P

    new_para = P(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para
