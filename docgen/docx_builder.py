from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from docgen.llm import LLMLogContext, generate_paragraph_sync, log_paragraph_disabled
from docgen.docx_merge import append_body_from_external_docx
from docgen.models import (
    DocumentConfig,
    FileParagraphElement,
    ImageElement,
    MathBlockElement,
    ParagraphElement,
    SectionNode,
    TableElement,
)
from docgen.word_utils import (
    BookmarkCounter,
    add_bookmark,
    add_hyperlink_to_bookmark,
    append_toc_entry_with_pageref,
)
from docgen.zip_bundle import ZipBundle


@dataclass
class RefTarget:
    bookmark: str
    label: str
    caption: str


def _json_slice(bundle: ZipBundle, key: str) -> dict[str, Any]:
    if not key:
        return {}
    raw = bundle.data.get(key)
    if raw is None:
        return {}
    if isinstance(raw, dict):
        return dict(raw)
    return {"value": raw}


def _basename_key(val: Any) -> str | None:
    if val is None:
        return None
    s = str(val).strip()
    if not s:
        return None
    return Path(s.replace("\\", "/")).name.lower()


def _section_bookmark_name(section_id: str) -> str:
    s = re.sub(r"[^A-Za-z0-9_]", "_", str(section_id).strip())
    if not s:
        s = "section"
    if s[0].isdigit():
        s = "s_" + s
    return "sec_" + s[:200]


def _collect_paths_in_order(config: DocumentConfig, bundle: ZipBundle) -> tuple[list[Path], list[Path]]:
    fig_paths: list[Path] = []
    tbl_paths: list[Path] = []

    def walk(node: SectionNode) -> None:
        for el in node.elements:
            if isinstance(el, TableElement):
                tbl_paths.append(bundle.resolve(el.csv_path))
            elif isinstance(el, ImageElement):
                fig_paths.append(bundle.resolve(el.img_path))
        for ch in node.children:
            walk(ch)

    for root in config.root_sections:
        walk(root)
    return fig_paths, tbl_paths


def _plan_refs(config: DocumentConfig, bundle: ZipBundle) -> tuple[list[RefTarget], list[RefTarget]]:
    fig_paths, tbl_paths = _collect_paths_in_order(config, bundle)
    figures = [
        RefTarget(bookmark=f"fig_{i + 1}", label=f"Figure {i + 1}", caption="")
        for i in range(len(fig_paths))
    ]
    tables = [
        RefTarget(bookmark=f"tbl_{i + 1}", label=f"Table {i + 1}", caption="")
        for i in range(len(tbl_paths))
    ]

    fi = 0
    ti = 0

    def fill_captions(node: SectionNode) -> None:
        nonlocal fi, ti
        for el in node.elements:
            if isinstance(el, TableElement):
                tables[ti].caption = el.caption or tables[ti].label
                ti += 1
            elif isinstance(el, ImageElement):
                figures[fi].caption = el.caption or figures[fi].label
                fi += 1
        for ch in node.children:
            fill_captions(ch)

    for root in config.root_sections:
        fill_captions(root)

    return figures, tables


def _ref_hints_from_json(
    ctx: dict[str, Any],
    fig_paths: list[Path],
    tbl_paths: list[Path],
    figures: list[RefTarget],
    tables: list[RefTarget],
) -> str:
    fk = _basename_key(ctx.get("figure") or ctx.get("image") or ctx.get("img"))
    tk = _basename_key(ctx.get("table") or ctx.get("csv"))
    lines: list[str] = []
    if fk:
        for i, p in enumerate(fig_paths):
            if p.name.lower() == fk:
                lines.append(f"Cite the image file {p.name} exactly as '{figures[i].label}' when needed.")
                break
    if tk:
        for i, p in enumerate(tbl_paths):
            if p.name.lower() == tk:
                lines.append(f"Cite the CSV {p.name} exactly as '{tables[i].label}' when needed.")
                break
    return "\n".join(lines)


def _split_figure_table_tokens(text: str) -> list[tuple[str, str | None]]:
    """Split text into (display_text, canonical_label_or_None). Canonical labels match fig_bm/tbl_bm keys."""
    pattern = re.compile(r"(?i)\b(figure\s+(\d+)|table\s+(\d+))\b")
    parts: list[tuple[str, str | None]] = []
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], None))
        if m.group(2) is not None:
            canon = f"Figure {m.group(2)}"
        else:
            canon = f"Table {m.group(3)}"
        parts.append((m.group(0), canon))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], None))
    return parts


def _render_text_with_refs(
    paragraph,
    text: str,
    fig_map: dict[str, str],
    tbl_map: dict[str, str],
    doc_part,
) -> None:
    label_to_bm = {**fig_map, **tbl_map}
    segments = _split_figure_table_tokens(text)
    if len(segments) == 1 and segments[0][1] is None:
        paragraph.add_run(text)
        return
    for seg, label in segments:
        if not seg:
            continue
        if label and label in label_to_bm:
            add_hyperlink_to_bookmark(paragraph, seg, label_to_bm[label], doc_part)
        else:
            paragraph.add_run(seg)


def _add_data_table(doc: Document, path: Path) -> None:
    df = pd.read_csv(path)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row.tolist()):
            cells[j].text = "" if pd.isna(val) else str(val)


def _add_picture(doc: Document, path: Path, width_inches: float = 5.2) -> None:
    doc.add_picture(str(path), width=Inches(width_inches))


def _add_text_blocks_from_file_body(
    doc: Document,
    text: str,
    fig_map: dict[str, str],
    tbl_map: dict[str, str],
    doc_part,
) -> None:
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    if not blocks and text.strip():
        blocks = [text.strip()]
    for b in blocks:
        joined = " ".join(line.strip() for line in b.splitlines() if line.strip())
        if not joined:
            continue
        p = doc.add_paragraph()
        _render_text_with_refs(p, joined, fig_map, tbl_map, doc_part)


def _add_front_matter(
    doc: Document,
    figures: list[RefTarget],
    tables: list[RefTarget],
    doc_part,
) -> None:
    doc.add_heading("List of Tables", level=1)
    for t in tables:
        p = doc.add_paragraph()
        append_toc_entry_with_pageref(p, doc_part, f"{t.label} — {t.caption}", t.bookmark)

    doc.add_heading("List of Figures", level=1)
    for f in figures:
        p = doc.add_paragraph()
        append_toc_entry_with_pageref(p, doc_part, f"{f.label} — {f.caption}", f.bookmark)


def build_docx(
    bundle: ZipBundle,
    config: DocumentConfig,
    out_path: Path,
    *,
    use_llm: bool = True,
    llm_log_ctx: LLMLogContext | None = None,
) -> None:
    figures, tables = _plan_refs(config, bundle)
    fig_paths, tbl_paths = _collect_paths_in_order(config, bundle)
    fig_bm = {f.label: f.bookmark for f in figures}
    tbl_bm = {t.label: t.bookmark for t in tables}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = title_p.add_run(config.title)
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph()

    _add_front_matter(doc, figures, tables, doc.part)
    doc.add_page_break()

    bm_counter = BookmarkCounter()
    fig_i = 0
    tbl_i = 0

    def render_section(node: SectionNode, depth: int) -> None:
        nonlocal fig_i, tbl_i
        heading = doc.add_heading(f"{node.section_id}  {node.name}", level=min(depth + 1, 9))
        add_bookmark(heading, _section_bookmark_name(node.section_id), bm_counter.next_id())

        for el in node.elements:
            if isinstance(el, ParagraphElement):
                ctx = _json_slice(bundle, el.json_key)
                hints = _ref_hints_from_json(ctx, fig_paths, tbl_paths, figures, tables)
                full_prompt = el.prompt.strip()
                if hints:
                    full_prompt = f"{full_prompt}\n\nCross-reference hints:\n{hints}"
                if use_llm:
                    body = generate_paragraph_sync(full_prompt, ctx, log_ctx=llm_log_ctx)
                else:
                    body = (
                        f"(LLM disabled) Facts: {ctx}. Original prompt: {el.prompt[:300]}"
                        + ("…" if len(el.prompt) > 300 else "")
                    )
                    if llm_log_ctx:
                        log_paragraph_disabled(
                            llm_log_ctx,
                            instruction_prompt=full_prompt,
                            facts_json=ctx,
                            body=body,
                        )
                p = doc.add_paragraph()
                _render_text_with_refs(p, body, fig_bm, tbl_bm, doc.part)

            elif isinstance(el, FileParagraphElement):
                fpath = bundle.resolve(el.text_path)
                raw_text = fpath.read_text(encoding="utf-8")
                _add_text_blocks_from_file_body(doc, raw_text, fig_bm, tbl_bm, doc.part)

            elif isinstance(el, TableElement):
                path = bundle.resolve(el.csv_path)
                cap = doc.add_paragraph()
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cap.add_run(f"{tables[tbl_i].label}: {tables[tbl_i].caption}")
                run.italic = True
                add_bookmark(cap, tables[tbl_i].bookmark, bm_counter.next_id())
                _add_data_table(doc, path)
                doc.add_paragraph()
                tbl_i += 1

            elif isinstance(el, ImageElement):
                path = bundle.resolve(el.img_path)
                cap = doc.add_paragraph()
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cap.add_run(f"{figures[fig_i].label}: {figures[fig_i].caption}")
                run.italic = True
                add_bookmark(cap, figures[fig_i].bookmark, bm_counter.next_id())
                _add_picture(doc, path)
                doc.add_paragraph()
                fig_i += 1

            elif isinstance(el, MathBlockElement):
                paths: list[Path] = []
                if el.docx_path.strip():
                    paths.append(bundle.resolve(el.docx_path.strip()))
                for rel in el.derivation_docx_paths:
                    rel = str(rel).strip()
                    if rel:
                        paths.append(bundle.resolve(rel))
                if not paths:
                    doc.add_paragraph("(math block: set docx_path / formula_docx and/or derivation_docx)")
                    continue
                if el.caption.strip():
                    cap = doc.add_paragraph()
                    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cap.add_run(el.caption.strip()).italic = True
                    doc.add_paragraph()
                for path in paths:
                    try:
                        append_body_from_external_docx(doc, path)
                    except (OSError, ValueError) as err:
                        err_p = doc.add_paragraph()
                        err_p.add_run(f"(could not merge math .docx {path.name}: {err})")

        for ch in node.children:
            render_section(ch, depth + 1)

    for root in config.root_sections:
        render_section(root, 0)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
